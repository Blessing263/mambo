"""Text extraction with provenance.

PDF: text per page via PyMuPDF. If a PDF is effectively scanned (very little
extractable text per page), fall back to OCR (Tesseract) by rasterising each page.
HTML: main-content extraction via trafilatura, with a BeautifulSoup fallback.
"""

from __future__ import annotations

import io

import fitz  # PyMuPDF
import pytesseract
import trafilatura
from bs4 import BeautifulSoup
from PIL import Image

# Below this average characters-per-page, treat the PDF as scanned and OCR it.
OCR_MIN_CHARS_PER_PAGE = 80
OCR_DPI = 200


def extract_pdf(content: bytes) -> dict:
    doc = fitz.open(stream=content, filetype="pdf")
    page_count = doc.page_count
    pages: list[tuple[int | None, str]] = []
    total_chars = 0
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        pages.append((i + 1, text))
        total_chars += len(text)

    ocr_used = False
    avg = total_chars / page_count if page_count else 0
    if page_count and avg < OCR_MIN_CHARS_PER_PAGE:
        ocr_used = True
        pages = []
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=OCR_DPI)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            pages.append((i + 1, pytesseract.image_to_string(img).strip()))

    title = (doc.metadata or {}).get("title") or None
    doc.close()
    return {
        "pages": pages,
        "title": title.strip() if title else None,
        "page_count": page_count,
        "ocr_used": ocr_used,
    }


def extract_html(content: bytes, url: str) -> dict:
    html = content.decode("utf-8", "ignore")
    text = trafilatura.extract(html, include_comments=False, include_tables=True) or ""
    soup = BeautifulSoup(html, "lxml")
    title = None
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    if not text:
        # Fallback: strip script/style then take visible text.
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(" ", strip=True)
    return {
        "pages": [(None, text)],
        "title": title,
        "page_count": None,
        "ocr_used": False,
    }


def extract_docx(content: bytes) -> dict:
    """Extract text from a .docx (paragraphs + tables) via python-docx."""
    import io
    from docx import Document  # noqa: PLC0415
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    title = None
    try:
        title = ((doc.core_properties.title or "")).strip() or None
    except Exception:
        pass
    return {
        "pages": [(None, text)],
        "title": title,
        "page_count": None,
        "ocr_used": False,
    }
